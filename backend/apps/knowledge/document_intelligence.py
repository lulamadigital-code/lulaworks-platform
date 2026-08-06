"""Document Intelligence — the single place the platform turns a document (or a
block of pasted text) into readable text and candidate line items.

Shared by the RFQ module and the Quotation module, and ready for Purchase Orders
and Contracts, so all of them behave identically and there is one implementation
to maintain rather than four that drift apart.

The line-item parsing itself is NOT reimplemented here: it delegates to
``apps.rfq.extraction``, which is deterministic-first and validated against real
Sibanye/Western Platinum documents. This module adds the two things the RFQ
parser did not need — reading text out of the many shapes a scope arrives in
(PDF, Word, Excel, an image, an email, a zip), and offering related items a
scope implies but did not spell out. Everything it returns is a *suggestion* for
a human to accept, edit or ignore; nothing here writes anything.
"""

import io
import json
import re
import zipfile
from email import message_from_bytes

from apps.rfq.extraction import parse_loose_lines, parse_rfq_text, to_decimal
from apps.rfq.extraction import extract_text as _pdf_text

#: What the upload control accepts. Extensions we can read to some degree; an
#: unknown type is not rejected, it just yields no text (and no items).
SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".csv", ".md", ".docx", ".xlsx", ".xls",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".eml", ".zip",
}


def _ext(name: str) -> str:
    name = (name or "").lower()
    return name[name.rfind("."):] if "." in name else ""


# ── Per-format text extractors (every optional library is lazy + graceful) ────

def _docx_text(data: bytes) -> str:
    try:
        import docx
    except ImportError:
        return ""
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception:
        return ""
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _xlsx_text(data: bytes) -> str:
    try:
        import openpyxl
    except ImportError:
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except Exception:
        return ""
    lines = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c not in (None, "")]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _image_text(data: bytes) -> str:
    """OCR an image (a photographed scope, a screenshot). Lazy — '' if the OCR
    toolchain is not installed, so a missing binary never 500s the request."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    try:
        return pytesseract.image_to_string(Image.open(io.BytesIO(data)))
    except Exception:
        return ""


def _eml_text(data: bytes) -> str:
    try:
        msg = message_from_bytes(data)
    except Exception:
        return ""
    parts = []
    for part in msg.walk() if msg.is_multipart() else [msg]:
        if part.get_content_type() == "text/plain":
            payload = part.get_payload(decode=True)
            if payload:
                parts.append(payload.decode(part.get_content_charset() or "utf-8",
                                             errors="ignore"))
    return "\n".join(parts)


def _plain_text(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def _zip_text(data: bytes) -> str:
    """Read every member we understand and concatenate — a zip of drawings and a
    BOQ becomes one body of text to read items from."""
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except Exception:
        return ""
    parts = []
    for info in zf.infolist():
        if info.is_dir() or _ext(info.filename) not in SUPPORTED_EXTENSIONS:
            continue
        if _ext(info.filename) == ".zip":      # do not recurse into nested zips
            continue
        try:
            parts.append(_text_from_bytes(info.filename, zf.read(info)))
        except Exception:
            continue
    return "\n".join(p for p in parts if p)


def _text_from_bytes(name: str, data: bytes) -> str:
    ext = _ext(name)
    if ext == ".pdf":
        return _pdf_text(io.BytesIO(data))          # pdfplumber text + OCR fallback
    if ext == ".docx":
        return _docx_text(data)
    if ext in (".xlsx", ".xls"):
        return _xlsx_text(data)
    if ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff"):
        return _image_text(data)
    if ext == ".eml":
        return _eml_text(data)
    if ext == ".zip":
        return _zip_text(data)
    if ext in (".txt", ".csv", ".md"):
        return _plain_text(data)
    # Unknown type: try to read it as text rather than fail outright.
    return _plain_text(data)


def extract_text_from_upload(uploaded_file) -> str:
    """Text out of an uploaded file, whatever its shape. Never raises on a bad
    file — a corrupt or unreadable upload yields '' and the caller simply finds
    no items."""
    try:
        data = uploaded_file.read()
    except Exception:
        return ""
    return _text_from_bytes(getattr(uploaded_file, "name", ""), data or b"")


# ── Candidate line items ──────────────────────────────────────────────────────

#: Units a job type prices in, used only to default a unit the text did not give.
_TYPE_UNIT = {"labour_hire": "hour", "plant_hire": "day",
              "project_management": "unit", "maintenance": "service"}


def extract_items(text: str, *, type_key: str | None = None,
                  company=None, user=None, use_ai: bool = False) -> list[dict]:
    """Candidate line items from text — the structured table first (rigid,
    high-confidence), then loose prose ("20 conveyor rollers"), de-duplicated by
    description. Returns plain dicts the estimator's grid fills in; the price is
    left blank unless the document actually stated one, because an invented price
    is worse than a blank.

    Deterministic-first. When ``use_ai`` and a provider is configured, Gemini is
    asked for the items the pattern parser missed (prose that is not in a tidy
    "N unit item" shape); it only *adds* descriptions, never overwrites, and
    never invents a price. AI is metered, so it runs on a deliberate action (a
    document upload), not on every keystroke."""
    if not text or not text.strip():
        return []

    lines = list(parse_rfq_text(text).lines) + list(parse_loose_lines(text))
    default_unit = _TYPE_UNIT.get(type_key or "", "each")

    items, seen = [], set()

    def _add(desc, qty, unit, price):
        key = (desc or "").strip().lower()
        if len(key) < 3 or key in seen:
            return
        seen.add(key)
        # The loose parser falls back to "each" when the text gave no unit; for a
        # labour or plant job that really means "unspecified", so the type
        # default wins. A unit the text stated explicitly is always kept.
        u = unit or default_unit
        if u == "each" and default_unit != "each":
            u = default_unit
        items.append({"description": desc.strip(), "qty": qty or "1",
                      "unit": u, "unit_price": price or ""})

    for ln in lines:
        _add(ln.description, f"{ln.qty:g}" if ln.qty else "1", ln.unit,
             f"{ln.unit_price:g}" if ln.unit_price else "")

    if use_ai:
        ai = _ai_json(company, user, _ITEMS_PROMPT, text, agent="quote_items")
        for ln in ai.get("lines", []) if isinstance(ai, dict) else []:
            price = ln.get("unit_price")
            _add(str(ln.get("description", "")),
                 f"{to_decimal(ln.get('qty', 1)) or 1:g}",
                 str(ln.get("unit", "") or ""),
                 f"{to_decimal(price):g}" if price not in (None, "", 0) else "")
    return items


# ── AI enrichment (Gemini via the metered gateway) ────────────────────────────
#
# Deterministic parsing is free and runs always; the model is the fallback for
# the shapes a regex cannot catch. Gated by ai_configured(), so with no key the
# whole platform still works — it simply extracts less.

_ITEMS_PROMPT = (
    "From this scope of work, list the quotable line items as strict JSON: "
    '{"lines":[{"description","qty","unit","unit_price"}]}. Use unit_price only '
    "if the text states a price; otherwise omit it. Do not invent items. "
    "Text:\n\n{text}"
)
_PO_PROMPT = (
    "Extract this customer purchase order as strict JSON with keys: po_number, "
    "po_date (YYYY-MM-DD), value (number), contact, site, payment_terms, "
    'delivery, and lines (array of {description, qty, unit, unit_price}). Use '
    "null for anything not present. Text:\n\n{text}"
)


def _ai_json(company, user, prompt_template: str, text: str, *, agent: str) -> dict:
    """Run one metered AI call and parse its JSON, tolerating ```json fences.
    Returns {} on any failure — a missing key, a provider error, unparseable
    output — so the caller always has its deterministic result to fall back to."""
    if not text.strip():
        return {}
    try:
        from apps.ai_platform.gateway import run_task
        from apps.ai_platform.providers import ai_configured
        from apps.ai_platform.routing import TaskType
    except ImportError:
        return {}
    if company is None or user is None or not ai_configured():
        return {}
    try:
        prompt = prompt_template.replace("{text}", text[:12000])
        # Document extraction routes to Gemini first, then fails over.
        resp = run_task(company, user, TaskType.EXTRACTION, prompt,
                        agent=agent, prompt_name=agent, json_mode=True)
        m = re.search(r"\{.*\}", resp.text, re.DOTALL)
        return json.loads(m.group(0)) if m else {}
    except Exception:
        return {}


#: Deterministic PO patterns beyond what parse_rfq_text already finds.
_VALUE_RE = re.compile(r"(?:total|order\s+value|po\s+value|grand\s+total)\s*:?\s*"
                       r"(R?\s?[\d][\d\s]*[.,]\d{2})", re.IGNORECASE)
_TERMS_RE = re.compile(r"(?:payment\s+terms?|terms)\s*:?\s*([^\n]{2,60})", re.IGNORECASE)
_DELIVERY_RE = re.compile(r"(?:delivery|deliver\s+to|ship\s+to)\s*:?\s*([^\n]{2,80})",
                          re.IGNORECASE)


def extract_po_fields(text: str, *, company=None, user=None,
                      use_ai: bool = False) -> dict:
    """Read a customer purchase order into the fields the PO form needs, so the
    estimator confirms rather than retypes. Deterministic-first (the same parser
    the RFQ front door uses), then Gemini fills what the patterns missed. Every
    value is a suggestion the user may correct before saving."""
    out = {"po_number": "", "po_date": "", "value": "", "contact": "",
           "site": "", "payment_terms": "", "delivery": "", "lines": []}
    if not text or not text.strip():
        return out

    fields = parse_rfq_text(text).fields
    if "po_number" in fields:
        out["po_number"] = fields["po_number"].value
    if "order_date" in fields:
        out["po_date"] = fields["order_date"].value.replace("/", "-")
    if "contact" in fields:
        out["contact"] = fields["contact"].value
    if "ship_to" in fields:
        out["site"] = fields["ship_to"].value
    if m := _VALUE_RE.search(text):
        out["value"] = f"{to_decimal(m.group(1)):.2f}"
    if m := _TERMS_RE.search(text):
        out["payment_terms"] = m.group(1).strip()
    if m := _DELIVERY_RE.search(text):
        out["delivery"] = m.group(1).strip()
    out["lines"] = extract_items(text)

    if use_ai:
        ai = _ai_json(company, user, _PO_PROMPT, text, agent="po_extraction")
        for key in ("po_number", "po_date", "value", "contact", "site",
                    "payment_terms", "delivery"):
            if not out[key] and ai.get(key):
                out[key] = str(ai[key])
        if not out["lines"] and isinstance(ai.get("lines"), list):
            out["lines"] = [
                {"description": str(x.get("description", "")),
                 "qty": f"{to_decimal(x.get('qty', 1)) or 1:g}",
                 "unit": str(x.get("unit", "") or "each"),
                 "unit_price": (f"{to_decimal(x.get('unit_price')):g}"
                                if x.get("unit_price") else "")}
                for x in ai["lines"] if x.get("description")
            ]
    return out


# ── Related-item suggestions (§7) ─────────────────────────────────────────────
#
# What a scope implies but rarely lists: a job that supplies rollers needs
# bearings, transport and consumables; installing anything needs labour. These
# are offered as chips, never added automatically, and never a price.

_RELATED = [
    (r"conveyor|roller|idler", ["Bearings", "Transport", "Consumables"]),
    (r"\binstall|installation|fit\b|mount", ["Installation labour"]),
    (r"pump|impeller|mechanical seal", ["Mechanical seals", "Gaskets", "Installation labour"]),
    (r"weld|fabricat|steel|structural", ["Welding consumables", "Transport", "Surface preparation"]),
    (r"paint|coat|corrosion", ["Surface preparation", "Consumables"]),
    (r"electric|cable|motor|panel", ["Cable & glands", "Commissioning", "Installation labour"]),
    (r"crane|lift|rigging|hoist", ["Rigging crew", "Mobilisation", "Transport"]),
    (r"pipe|pipeline|flange|valve", ["Gaskets", "Fasteners", "Installation labour"]),
    (r"maintenance|service|inspection", ["Consumables", "Report & sign-off"]),
]


def suggest_related_items(text: str, existing: list[str] | None = None) -> list[str]:
    """Related items a scope implies. De-duplicated, and never anything the
    estimator already has on the quotation."""
    import re

    if not text or not text.strip():
        return []
    have = {e.strip().lower() for e in (existing or [])}
    out, seen = [], set()
    for pattern, related in _RELATED:
        if not re.search(pattern, text, re.IGNORECASE):
            continue
        for item in related:
            key = item.lower()
            if key in have or key in seen:
                continue
            seen.add(key)
            out.append(item)
    return out[:6]
